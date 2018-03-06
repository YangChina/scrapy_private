from scrapy import Request
from scrapy.spiders import Spider
from scrapy_project.items import Exam_Gongwuyuan_shenlun_Item
import os
import re
import time
import random

from docx import Document
from docx.oxml.ns import qn

class ShenlunSpider(Spider):
    name = 'shenlun'
    
    start_urls = {
        'http://www.gwyzk.com/e/action/ListInfo/?classid=477&page=1',
    }
    custom_settings = {
        'ITEM_PIPELINES':{},
        'DOWNLOAD_DELAY' :'2',
    }

    basePath = './spider_data/shenlunfanwen/'

    def parse(self, response):
        #print('ShenlunSpider -->> parse -->> response.body:::',response.body.decode('gbk'))
        article_lists = response.xpath('//div[@class="list"]//li')
        
        for i, article in enumerate(article_lists):
            #print('ShenlunSpider -->> parse -->> article :::',article.xpath('.//a/@href').extract()[0])
            #print('ShenlunSpider -->> parse -->> article :::',article.xpath('.//a/text()').extract())
            
            if type(article.xpath('.//a/@href').extract()) != type([]):
                print('ShenlunSpider -->> parse -->> article_url :::','抓取URL不是list......')
                continue
            else:
                article_url = article.xpath('.//a/@href').extract()[0]
            
            if article.xpath('.//a/text()').extract():
                article_title = article.xpath('.//a/text()').extract()[0]
                article_path = basePath + article_title + '.docx'
            else:
                print('ShenlunSpider -->> parse -->> article_title :::','标题非法，跳过......')
            
            if os.path.exists(article_path):
                print('ShenlunSpider -->> parse -->> os.path.exists :::',article_title,' 文件已存在，跳过......')
                continue
            
            if article_url.find('http://www.gwyzk.com/html/') < 0:
                print('ShenlunSpider -->> parse -->> article_url :::','不是有效的URL，跳过......')
                continue
            #print('ShenlunSpider -->> parse -->> article_url :::',article_url)
            #print('ShenlunSpider -->> parse -->> article_title :::',article_title)
            yield Request(article_url,callback=self.parse_article_content)

        
        pages = response.xpath('//div[@class="fenye"]//a')
        next_url = ''
        for i, page_tag in enumerate(pages):
            page = page_tag.xpath('./text()').extract()[0]
            page_url = page_tag.xpath('./@href').extract()
            if page == '下一页':
                next_url = page_url[0]
                break
            #print('ShenlunSpider -->> parse -->> page :::',page)

        print('ShenlunSpider -->> parse -->> next_url :::',next_url)

        if next_url:
            yield Request(next_url,callback=self.parse)
 
    def parse_article_content(self,response):
        time.sleep(random.randint(2,10))
        
        doc = Document()
        doc.styles['Normal'].font.name = u'华文仿宋'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文仿宋')

        article_title = response.xpath('//div[@class="content"]/h1/text()').extract()[0]
        article_contents = response.xpath('//div[@class="zhengwen ptzw"]//p').extract()

        print('ShenlunSpider -->> parse_article_content -->> article_title :::',article_title)
        #print('ShenlunSpider -->> parse_article_content -->> article_contents :::',article_contents)

        doc.add_heading(article_title, 0)

        dr = re.compile(r'<[^>]+>',re.S)

        for i, article_content_p in enumerate(article_contents):
            article_content_p = dr.sub('',article_content_p)
            if article_content_p.find('相同主题') >= 0 or \
               article_content_p.find('相关主题') >= 0 or \
               article_content_p.find('相关文章') >= 0 or \
               article_content_p.find('欢迎关注') >= 0:
                break
            
            doc.add_paragraph(article_content_p)
            #print('ShenlunSpider -->> parse_article_content -->> article_content_p :::',article_content_p)
        
        doc.save(basePath + article_title + '.docx')
            
