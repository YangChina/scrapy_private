from scrapy import Request
from scrapy.spiders import Spider
from scrapy_project.items import Exam_Gongwuyuan_shenlun_Item
import os
import re
import time
import random
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from scrapy.selector import Selector 
from scrapy.http import FormRequest

class ShenlunSpider(Spider):
    print('================ test no. 2 =====================')
    name = 'chinagwy_tiku'
    
    start_urls = {
        'http://exam.chinagwy.org/index.php?mod=xingce&act=zhenti',
    }

    custom_settings = {
        'ITEM_PIPELINES':{},
        'DOWNLOAD_DELAY' :'0',
    }

    headers = {
        "Host":"exam.chinagwy.org",
        "Accept":"text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language":"zh-CN,zh;q=0.8",
        "Accept-Encoding":"gzip, deflate, sdch",
        "Referer":"http://exam.chinagwy.org/index.php",
        "Cookie":"Hm_lvt_7c630b740ceaad93a6ccf6dc79182b00=1520258423; Hm_lpvt_7c630b740ceaad93a6ccf6dc79182b00=1520258423; PHPSESSID=g5r5nk1ib3smvlumin8pmcid56",
        "Connection":"keep-alive",
        "Proxy-Connection":"keep-alive",
        "Upgrade-Insecure-Requests":"1",
        "User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/55.0.2883.87 Safari/537.36"
    }

    def parse(self, response):
        exam_lists = Selector(text=response.body).xpath('//td[@class="enter_bt"]')
        #print('chinagwy_tiku -->> parse -->> exam_lists :::',exam_lists)
        #print(type(exam_lists[0]))
        for i, exam in enumerate(exam_lists):
            if i >= 1:
                break
            exam_url = 'http://exam.chinagwy.org/index.php' + exam.xpath('./a/@href').extract()[0]
            #print('chinagwy_tiku -->> parse -->> exam :::',exam)
            print('chinagwy_tiku -->> parse -->> exam_url :::',exam_url)
            yield Request(exam_url, headers=self.headers, callback=self.parse_exam)
            '''
            if type(article.xpath('.//a/@href').extract()) != type([]):
                print('ShenlunSpider -->> parse -->> article_url :::','抓取URL不是list......')
                continue
            else:
                article_url = article.xpath('.//a/@href').extract()[0]
            
            if article.xpath('.//a/text()').extract():
                article_title = article.xpath('.//a/text()').extract()[0]
                article_path = article_title + '.docx'
            else:
                print('ShenlunSpider -->> parse -->> article_title :::','标题非法，跳过......')
            
            if os.path.exists(article_path):
                print('ShenlunSpider -->> parse -->> os.path.exists :::',article_title,' 文件已存在，跳过......')
                continue
            
            if article_url.find('http://www.gwyzk.com/html/') < 0:
                print('ShenlunSpider -->> parse -->> article_url :::','不是有效的URL，跳过......')
                continue
            #print('ShenlunSpider -->> parse -->> article_url :::',article_url)
            #print('ShenlunSpider -->> parse -->> article_title :::',article_title)
            yield Request(article_url,callback=self.parse_article_content)

        
        pages = response.xpath('//div[@class="fenye"]//a')
        next_url = ''
        for i, page_tag in enumerate(pages):
            page = page_tag.xpath('./text()').extract()[0]
            page_url = page_tag.xpath('./@href').extract()
            if page == '下一页':
                next_url = page_url[0]
                break
            #print('ShenlunSpider -->> parse -->> page :::',page)

        print('ShenlunSpider -->> parse -->> next_url :::',next_url)

        if next_url:
            yield Request(next_url,callback=self.parse)
 '''
    def parse_exam(self,response):
        '''
        1.交卷。
        Request URL:http://exam.chinagwy.org/index.php?mod=exercise&act=submit
        Request Method:POST
        FormData
            paper_id:9929449
            flag:2
        2.response 查看报告
        http://exam.chinagwy.org/index.php?mod=exercise&act=report&tid=9929449
        3.查看解析
        http://exam.chinagwy.org/index.php?mod=exercise&act=solution&tid=9929449

        '''


        print('chinagwy_tiku -->> parse_exam -->> response :::',response)
        
        if response.url.find('act=paper') >= 0:
            # 载入试题页面
            # 发送提交请求

            pass
        
        if response.url.find('act=report') >= 0:
            # 载入报告页面
            # 发送查看解析请求
            
            pass
        
        if response.url.find('act=paper') >= 0:
            exam_url = response.url.replace('act=paper', 'act=solution')
            print('chinagwy_tiku -->> parse_exam -->> exam_url.replace :::',exam_url)
            yield Request(exam_url,headers=self.headers,callback=self.parse_exam)
            return
            
        print('======================================')
        print(response.body.decode('utf8'))



        exam_title = Selector(text=response.body.decode('utf8')).xpath('//div[@class="hc_title"]/text()')
        print('chinagwy_tiku -->> parse_exam -->> exam_title :::',exam_title)
        exam_pary_cnt = len(Selector(text=response.body).xpath('//ul[@class="hc_tab"]//li'))
        print('chinagwy_tiku -->> parse_exam -->> exam_pary_cnt :::',exam_pary_cnt)

        exam_content = Selector(text=response.body).xpath('//td[@class="enter_bt"]')
        #time.sleep(random.randint(2,10))
        print('chinagwy_tiku -->> parse_exam -->> response :::',response)
        #print('chinagwy_tiku -->> parse_exam -->> response.body :::',response.body.decode('utf8'))

        print('chinagwy_tiku -->> parse_exam -->> response.url :::',response.url.encode('utf8'))
        exam_id = response.url[response.url.find('tid=')+4:]
        
        print('chinagwy_tiku -->> parse_exam -->> exam_id :::',exam_id)
        res = response.body.decode('utf8')
        
        requests = []
        for idx in range(exam_pary_cnt):
            print(idx)
            part_title = Selector(text=response.body).xpath('//ul[@class="hc_tab"]//li')[idx].xpath('./a/text()').extract()[0]
            #print('chinagwy_tiku -->> parse_exam -->> part_title :::',part_title)
            page_url = 'http://exam.chinagwy.org/index.php?mod=exercise&act=paper'
            page_form = {
                "chapter_index":str(idx),
                "paper_id":exam_id,
            }
            yield FormRequest(page_url,headers=self.headers,meta={'exam_title':exam_title,'part_title':part_title},callback=self.parse_exam_part,formdata=page_form)
            
            

        '''
        
        requests.append(request)

        
        if res.find('用户登录') >= 0:
            return Request(exam_url,headers=self.headers,callback=self.parse_login)
        

        
        doc = Document()
        doc.styles['Normal'].font.name = u'华文仿宋'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文仿宋')

        article_title = response.xpath('//div[@class="content"]/h1/text()').extract()[0]
        article_contents = response.xpath('//div[@class="zhengwen ptzw"]//p').extract()

        print('ShenlunSpider -->> parse_article_content -->> article_title :::',article_title)
        #print('ShenlunSpider -->> parse_article_content -->> article_contents :::',article_contents)

        doc.add_heading(article_title, 0)

        dr = re.compile(r'<[^>]+>',re.S)

        for i, article_content_p in enumerate(article_contents):
            article_content_p = dr.sub('',article_content_p)
            if article_content_p.find('相同主题') >= 0 or \
               article_content_p.find('相关主题') >= 0 or \
               article_content_p.find('相关文章') >= 0 or \
               article_content_p.find('欢迎关注') >= 0:
                break
            
            doc.add_paragraph(article_content_p)
            #print('ShenlunSpider -->> parse_article_content -->> article_content_p :::',article_content_p)
        
        doc.save(article_title+'.docx')
        '''
    def parse_exam_part(self,response):
        print('chinagwy_tiku -->> parse_exam_part -->> response :::',response)
        #print('chinagwy_tiku -->> parse_exam_part -->> response.body :::',response.body.decode('utf8'))

        '''
        doc = Document()
        doc.styles['Normal'].font.name = u'华文仿宋'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文仿宋')
        '''
        exam_title = response.meta['exam_title']
        print('chinagwy_tiku -->> parse_exam_part -->> exam_title :::',exam_title)
        part_title = response.meta['part_title']
        print('chinagwy_tiku -->> parse_exam_part -->> part_title :::',part_title)
        exam_content = response.body.decode('utf8')
        #print('chinagwy_tiku -->> parse_exam_part -->> exam_content :::',exam_content)
        '''
        doc.add_heading(article_title, 0)

        dr = re.compile(r'<[^>]+>',re.S)

        for i, article_content_p in enumerate(article_contents):
            article_content_p = dr.sub('',article_content_p)
            if article_content_p.find('相同主题') >= 0 or \
               article_content_p.find('相关主题') >= 0 or \
               article_content_p.find('相关文章') >= 0 or \
               article_content_p.find('欢迎关注') >= 0:
                break
            
            doc.add_paragraph(article_content_p)
            #print('ShenlunSpider -->> parse_article_content -->> article_content_p :::',article_content_p)
        
        doc.save(basePath + article_title + '.docx')'''
            